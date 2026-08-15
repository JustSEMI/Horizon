import sys
import os
import argparse
import time
import pandas as pd
from playwright.sync_api import sync_playwright

def main():
    parser = argparse.ArgumentParser(description="Google Maps Scraper")
    parser.add_argument("--query", type=str, required=True, help="Search query")
    parser.add_argument("--max", type=int, default=20, help="Maximum results to scrape")
    parser.add_argument("--output", type=str, default="output.xlsx", help="Output file path")
    parser.add_argument("--format", type=str, default="xlsx", help="Export format (xlsx, docx, pdf, html)")
    args = parser.parse_args()

    query = args.query
    max_results = args.max
    output_path = args.output

    print(f"[INFO] Starting Playwright Scraper for query: '{query}'...")
    print(f"[INFO] Max results to fetch: {max_results}")

    results = []

    try:
        with sync_playwright() as p:
            print("[INFO] Launching Chromium browser (headless mode)...")
            browser = p.chromium.launch(headless=True)
            context = browser.new_context(locale="en-US")
            page = context.new_page()

            print("[INFO] Navigating to Google Maps...")
            page.goto("https://www.google.com/maps", timeout=60000)
            
            print("[INFO] Searching...")
            page.wait_for_selector('input[name="q"]', timeout=15000)
            page.fill('input[name="q"]', query)
            page.keyboard.press("Enter")

            try:
                page.wait_for_selector('div[role="feed"]', timeout=15000)
            except Exception as e:
                print(f"[ERROR] Could not find results feed. Maybe no results found or slow internet.")
                browser.close()
                sys.exit(1)

            print("[INFO] Scraping results... This may take a while as we scroll down.")

            previously_counted = 0
            retries = 0
            while True:
                listings = page.locator('div[role="feed"] > div > div > a').all()
                current_count = len(listings)
                
                if current_count >= max_results:
                    break
                    
                if current_count == previously_counted:
                    retries += 1
                    if retries > 5:
                        print("[INFO] Reached the end of the list or no more results loading.")
                        break
                else:
                    retries = 0
                
                previously_counted = current_count
                
                if current_count > 0:
                    try:
                        listings[-1].hover(timeout=1000)
                    except:
                        pass
                
                # Force scroll feed element directly using JS
                try:
                    page.evaluate('''(document.querySelector("div[role=\'feed\']") || document.querySelector("div.m6QErb.DxyBCb")).scrollBy(0, 5000)''')
                except:
                    pass
                    
                page.mouse.wheel(0, 5000)
                time.sleep(2.5)

            listings = page.locator('div[role="feed"] > div > div > a').all()
            total_to_process = min(len(listings), max_results)
            print(f"[INFO] Found {len(listings)} listings. Processing {total_to_process}...")

            for i in range(total_to_process):
                print(f"[INFO] Extracting place {i+1}/{total_to_process}...")
                listing = listings[i]
                
                try:
                    listing.scroll_into_view_if_needed()
                    listing.click()
                    page.wait_for_selector('h1', timeout=10000)
                    time.sleep(1.5)
                    
                    name = ""
                    try:
                        name = page.locator('h1').inner_text()
                    except:
                        pass
                    if not name:
                        try:
                            # Fallback: get name from aria-label of the listing card
                            name = listing.get_attribute('aria-label') or ""
                        except:
                            pass
                    
                    rating = ""
                    try:
                        rating_text = page.locator('div.F7nice').inner_text()
                        parts = rating_text.split('\\n')
                        if len(parts) > 0:
                            rating = parts[0]
                    except:
                        pass
                        
                    category = ""
                    try:
                        category = page.locator('button.DkEaL').inner_text()
                    except:
                        pass
                        
                    address = ""
                    phone = ""
                    website = ""
                    raw_url = page.url
                    # Clean the URL
                    clean_url = raw_url
                    if "/@" in raw_url:
                        clean_url = raw_url.split("/@")[0]
                    elif "/data=" in raw_url:
                        clean_url = raw_url.split("/data=")[0]
                    
                    info_items = page.locator('button[data-item-id]').all()
                    for item in info_items:
                        try:
                            aria = item.get_attribute('aria-label') or ""
                            if "Address:" in aria or "Alamat:" in aria:
                                address = aria.replace("Address:", "").replace("Alamat:", "").strip()
                            elif "Phone:" in aria or "Telepon:" in aria:
                                phone = aria.replace("Phone:", "").replace("Telepon:", "").strip()
                            elif "Website:" in aria:
                                website = aria.replace("Website:", "").strip()
                        except:
                            pass
                            
                    results.append({
                        "Name": name,
                        "Category": category,
                        "Rating": rating,
                        "Address": address,
                        "Phone": phone,
                        "Website": website,
                        "URL": clean_url
                    })
                except Exception as e:
                    print(f"[WARNING] Failed to extract data for item {i+1}: {str(e)}")

            browser.close()
            
    except Exception as e:
        print(f"[ERROR] Playwright execution failed: {e}")

    if results:
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            df = pd.DataFrame(results)
            fmt = args.format.lower()
            
            if fmt == "xlsx":
                df.to_excel(output_path, index=False)
                
            elif fmt == "html":
                html = df.to_html(index=False, render_links=True, classes="table table-striped")
                styled_html = f"<html><head><style>table {{font-family: Arial, sans-serif; border-collapse: collapse; width: 100%;}} td, th {{border: 1px solid #ddd; padding: 8px;}} tr:nth-child(even){{background-color: #f2f2f2;}} th {{padding-top: 12px; padding-bottom: 12px; text-align: left; background-color: #04AA6D; color: white;}}</style></head><body><h2>Google Maps Scraper Results</h2>{html}</body></html>"
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(styled_html)
                    
            elif fmt == "docx":
                try:
                    import docx
                    doc = docx.Document()
                    doc.add_heading('Google Maps Scraper Results', 0)
                    
                    # Create table
                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.style = 'Table Grid'
                    
                    # Add headers
                    hdr_cells = table.rows[0].cells
                    for i, col_name in enumerate(df.columns):
                        hdr_cells[i].text = str(col_name)
                        
                    # Add rows
                    for index, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for i, val in enumerate(row):
                            row_cells[i].text = str(val) if val else ""
                            
                    doc.save(output_path)
                except ImportError:
                    print("[ERROR] python-docx library is required for Word export.")
                    
            elif fmt == "pdf":
                try:
                    from fpdf import FPDF
                    
                    class PDF(FPDF):
                        def header(self):
                            self.set_font("helvetica", "B", 15)
                            self.cell(w=0, h=10, text="Google Maps Scraper Results", border=0, new_x="LMARGIN", new_y="NEXT", align="C")
                            
                    pdf = PDF()
                    pdf.add_page()
                    pdf.set_font("helvetica", size=8)
                    
                    # For PDF, table rendering can be complex, so we will list items
                    for index, row in df.iterrows():
                        pdf.set_font("helvetica", "B", 10)
                        
                        name = str(row['Name']).encode('latin-1', 'replace').decode('latin-1')
                        pdf.cell(w=0, h=8, text=f"{index+1}. {name}", border=0, new_x="LMARGIN", new_y="NEXT", align='L')
                        pdf.set_font("helvetica", "", 8)
                        
                        cat = str(row.get('Category', '')).encode('latin-1', 'replace').decode('latin-1')
                        rating = str(row.get('Rating', '')).encode('latin-1', 'replace').decode('latin-1')
                        address = str(row.get('Address', '')).encode('latin-1', 'replace').decode('latin-1')
                        phone = str(row.get('Phone', '')).encode('latin-1', 'replace').decode('latin-1')
                        website = str(row.get('Website', '')).encode('latin-1', 'replace').decode('latin-1')
                        url = str(row.get('URL', ''))
                        
                        import textwrap
                        
                        cat_text = f"Category: {cat} | Rating: {rating}"
                        addr_text = f"Address: {address}"
                        contact_text = f"Phone: {phone} | Web: {website}"
                        
                        for line in textwrap.wrap(cat_text, width=110, break_long_words=True):
                            pdf.cell(w=0, h=5, text=line, border=0, new_x="LMARGIN", new_y="NEXT", align="L")
                            
                        for line in textwrap.wrap(addr_text, width=110, break_long_words=True):
                            pdf.cell(w=0, h=5, text=line, border=0, new_x="LMARGIN", new_y="NEXT", align="L")
                            
                        for line in textwrap.wrap(contact_text, width=110, break_long_words=True):
                            pdf.cell(w=0, h=5, text=line, border=0, new_x="LMARGIN", new_y="NEXT", align="L")
                        
                        if url:
                            pdf.set_text_color(0, 0, 255)
                            pdf.cell(w=0, h=5, text="View on Google Maps", border=0, new_x="LMARGIN", new_y="NEXT", link=url)
                            pdf.set_text_color(0, 0, 0)
                            
                        pdf.ln(3)
                        
                    pdf.output(output_path)
                except ImportError:
                    print("[ERROR] fpdf2 library is required for PDF export.")
            
            print(f"[SUCCESS] Scraped {len(results)} items and saved to {output_path}")
        except Exception as e:
            print(f"[ERROR] Failed to save {args.format.upper()} file: {e}")
    else:
        print("[INFO] No results found to save.")

if __name__ == "__main__":
    main()
